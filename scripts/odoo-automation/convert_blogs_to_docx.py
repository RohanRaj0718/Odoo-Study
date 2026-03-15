#!/usr/bin/env python3
"""
Convert all 14 blog Markdown files to Word (.docx) format.
Handles: headings, bold, italic, tables, bullet lists, numbered lists,
         code blocks, inline code, blockquotes, horizontal rules, links.
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── File list ──────────────────────────────────────────────────────────
BASE = r"C:\Odoo Study"

BLOG_FILES = [
    "BLOG_Subcontracting_Process_Odoo_19.md",
    "BLOG_Subcontracting_V2_Tutorial.md",
    "BLOG_Subcontracting_V3_Storytelling.md",
    "BLOG_Subcontracting_V4_Listicle.md",
    "BLOG_Subcontracting_V5_Technical.md",
    "BLOG_Subcontracting_V6_Business.md",
    "BLOG_Subcontracting_V7_DocStyle.md",
    "BLOG_Work_Centers_Routing_Odoo.md",
    "BLOG_Work_Centers_V2_Tutorial.md",
    "BLOG_Work_Centers_V3_Storytelling.md",
    "BLOG_Work_Centers_V4_Listicle.md",
    "BLOG_Work_Centers_V5_Technical.md",
    "BLOG_Work_Centers_V6_Business.md",
    "BLOG_Work_Centers_V7_DocStyle.md",
]

# ── Helpers ────────────────────────────────────────────────────────────

def setup_styles(doc):
    """Configure document styles for consistent formatting."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15

    # Heading styles
    heading_configs = {
        'Heading 1': (Pt(22), RGBColor(0x1A, 0x1A, 0x2E), True, Pt(18), Pt(8)),
        'Heading 2': (Pt(17), RGBColor(0x2C, 0x3E, 0x50), True, Pt(14), Pt(6)),
        'Heading 3': (Pt(14), RGBColor(0x34, 0x49, 0x5E), True, Pt(10), Pt(4)),
        'Heading 4': (Pt(12), RGBColor(0x5D, 0x6D, 0x7E), True, Pt(8), Pt(4)),
    }
    for name, (size, color, bold, sp_before, sp_after) in heading_configs.items():
        if name in doc.styles:
            s = doc.styles[name]
            s.font.name = 'Calibri'
            s.font.size = size
            s.font.color.rgb = color
            s.font.bold = bold
            s.paragraph_format.space_before = sp_before
            s.paragraph_format.space_after = sp_after


def add_formatted_text(paragraph, text):
    """Parse inline markdown (bold, italic, code, links) and add runs."""
    # Pattern order matters: bold-italic first, then bold, italic, inline code, links
    patterns = [
        (r'\*\*\*(.*?)\*\*\*', 'bold_italic'),    # ***bold italic***
        (r'\*\*(.*?)\*\*', 'bold'),                 # **bold**
        (r'__(.*?)__', 'bold'),                      # __bold__
        (r'\*(.*?)\*', 'italic'),                    # *italic*
        (r'_(.*?)_', 'italic'),                      # _italic_
        (r'`(.*?)`', 'code'),                        # `code`
        (r'\[(.*?)\]\((.*?)\)', 'link'),             # [text](url)
    ]
    
    # Build a combined pattern
    combined = '|'.join(f'(?P<g{i}>{p})' for i, (p, _) in enumerate(patterns))
    
    pos = 0
    for m in re.finditer(combined, text):
        # Add plain text before this match
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        
        # Determine which pattern matched
        matched_text = m.group()
        
        if re.match(r'\*\*\*(.*?)\*\*\*', matched_text):
            inner = re.match(r'\*\*\*(.*?)\*\*\*', matched_text).group(1)
            run = paragraph.add_run(inner)
            run.bold = True
            run.italic = True
        elif re.match(r'\*\*(.*?)\*\*', matched_text):
            inner = re.match(r'\*\*(.*?)\*\*', matched_text).group(1)
            run = paragraph.add_run(inner)
            run.bold = True
        elif re.match(r'__(.*?)__', matched_text):
            inner = re.match(r'__(.*?)__', matched_text).group(1)
            run = paragraph.add_run(inner)
            run.bold = True
        elif re.match(r'\*(.*?)\*', matched_text):
            inner = re.match(r'\*(.*?)\*', matched_text).group(1)
            run = paragraph.add_run(inner)
            run.italic = True
        elif re.match(r'_(.*?)_(?!_)', matched_text):
            inner = re.match(r'_(.*?)_', matched_text).group(1)
            run = paragraph.add_run(inner)
            run.italic = True
        elif re.match(r'`(.*?)`', matched_text):
            inner = re.match(r'`(.*?)`', matched_text).group(1)
            run = paragraph.add_run(inner)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
            # Add light background shading
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
            run._element.get_or_add_rPr().append(shading_elm)
        elif re.match(r'\[(.*?)\]\((.*?)\)', matched_text):
            link_match = re.match(r'\[(.*?)\]\((.*?)\)', matched_text)
            link_text = link_match.group(1)
            run = paragraph.add_run(link_text)
            run.font.color.rgb = RGBColor(0x21, 0x96, 0xF3)
            run.underline = True
        
        pos = m.end()
    
    # Add remaining plain text
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_table(doc, header_row, data_rows):
    """Add a formatted table to the document."""
    num_cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    for i, cell_text in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(cell_text.strip())
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Dark header background
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2C3E50" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Data rows
    for row_idx, row_data in enumerate(data_rows):
        for col_idx in range(min(len(row_data), num_cols)):
            cell = table.rows[1 + row_idx].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            add_formatted_text(p, row_data[col_idx].strip())
            p.style = doc.styles['Normal']
            for run in p.runs:
                run.font.size = Pt(10)
            # Alternate row shading
            if row_idx % 2 == 0:
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8F9FA" w:val="clear"/>')
                cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Add spacing after table
    doc.add_paragraph()


def parse_table_lines(lines):
    """Parse markdown table lines into header and data rows."""
    if len(lines) < 2:
        return None, None
    
    def split_row(line):
        # Remove leading/trailing pipes and split
        cells = line.strip().strip('|').split('|')
        return [c.strip() for c in cells]
    
    header = split_row(lines[0])
    # Skip separator line (lines[1] contains ---)
    data = []
    for line in lines[2:]:
        if line.strip():
            data.append(split_row(line))
    
    return header, data


def convert_md_to_docx(md_path, docx_path):
    """Convert a Markdown file to a Word document."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    doc = Document()
    setup_styles(doc)
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    i = 0
    in_code_block = False
    code_lines = []
    code_lang = ''
    
    while i < len(lines):
        line = lines[i]
        
        # ── Code block handling ──
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lang = line.strip()[3:].strip()
                code_lines = []
                i += 1
                continue
            else:
                # End of code block - output it
                in_code_block = False
                code_text = '\n'.join(code_lines)
                if code_text.strip():
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
                    # Background shading for entire paragraph
                    pPr = p._element.get_or_add_pPr()
                    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
                    pPr.append(shading_elm)
                i += 1
                continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # ── Blank line ──
        if not line.strip():
            i += 1
            continue
        
        # ── Table detection ──
        if '|' in line and i + 1 < len(lines) and re.match(r'^\s*\|?\s*[-:]+\s*\|', lines[i + 1]):
            table_lines = []
            j = i
            while j < len(lines) and '|' in lines[j] and lines[j].strip():
                table_lines.append(lines[j])
                j += 1
            header, data = parse_table_lines(table_lines)
            if header and data:
                add_table(doc, header, data)
            i = j
            continue
        
        # ── Horizontal rule ──
        if re.match(r'^(\*{3,}|-{3,}|_{3,})\s*$', line.strip()):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add a horizontal line via border
            pPr = p._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue
        
        # ── Headings ──
        heading_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            # Remove trailing # marks
            heading_text = re.sub(r'\s*#+\s*$', '', heading_text)
            style_name = f'Heading {min(level, 4)}'
            p = doc.add_paragraph(style=style_name)
            add_formatted_text(p, heading_text)
            i += 1
            continue
        
        # ── Blockquote ──
        if line.strip().startswith('>'):
            quote_text = re.sub(r'^>\s*', '', line.strip())
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add left border
            pPr = p._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:left w:val="single" w:sz="12" w:space="8" w:color="714B67"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            add_formatted_text(p, quote_text)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue
        
        # ── Bullet list (-, *, +) ──
        bullet_match = re.match(r'^(\s*)([-*+])\s+(.*)', line)
        if bullet_match:
            indent_level = len(bullet_match.group(1)) // 2
            bullet_text = bullet_match.group(3)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(1.27 + indent_level * 0.63)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_formatted_text(p, bullet_text)
            i += 1
            continue
        
        # ── Numbered list ──
        num_match = re.match(r'^(\s*)(\d+)[.)]\s+(.*)', line)
        if num_match:
            indent_level = len(num_match.group(1)) // 2
            item_text = num_match.group(3)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Cm(1.27 + indent_level * 0.63)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_formatted_text(p, item_text)
            i += 1
            continue
        
        # ── Screenshot / Image placeholder ──
        img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line.strip())
        if img_match:
            alt_text = img_match.group(1)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(f'[Image: {alt_text}]')
            run.italic = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.font.size = Pt(10)
            i += 1
            continue
        
        # ── Regular paragraph ──
        p = doc.add_paragraph()
        add_formatted_text(p, line.strip())
        i += 1
    
    doc.save(docx_path)
    return docx_path


# ── Main ───────────────────────────────────────────────────────────────

def main():
    print("=" * 60)
    print("  Blog MD → DOCX Converter")
    print("=" * 60)
    
    success = 0
    failed = 0
    
    for md_file in BLOG_FILES:
        md_path = os.path.join(BASE, md_file)
        docx_file = md_file.replace('.md', '.docx')
        docx_path = os.path.join(BASE, docx_file)
        
        if not os.path.exists(md_path):
            print(f"  ✗ NOT FOUND: {md_file}")
            failed += 1
            continue
        
        try:
            convert_md_to_docx(md_path, docx_path)
            # Verify file was created and has content
            size = os.path.getsize(docx_path)
            print(f"  ✓ {docx_file} ({size:,} bytes)")
            success += 1
        except Exception as e:
            print(f"  ✗ FAILED: {md_file} → {e}")
            failed += 1
    
    print(f"\n{'=' * 60}")
    print(f"  Results: {success} converted, {failed} failed")
    print(f"  Output folder: {BASE}")
    print(f"{'=' * 60}")


if __name__ == '__main__':
    main()
