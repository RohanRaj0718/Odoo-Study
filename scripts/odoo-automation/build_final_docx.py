"""Create final enhanced client operations guide with embedded screenshots.
Follows the style of Profit and Loss Reporting.docx and Account Grouping Doc.docx."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# ── STYLES ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

def add_nav_path(text):
    nav = doc.add_paragraph()
    nav.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = nav.add_run(text)
    r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    r.font.size = Pt(10)

def add_image(img_path, caption, width=Inches(6.3)):
    """Insert image centered with caption below (matching reference doc pattern)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=width)
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def add_bold_para(bold_text, normal_text=''):
    p = doc.add_paragraph()
    r = p.add_run(bold_text)
    r.bold = True
    if normal_text:
        p.add_run(normal_text)

def add_tip(text):
    p = doc.add_paragraph()
    r = p.add_run('Tip: ')
    r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x7D, 0x1F)
    p.add_run(text)

# ══════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════
title = doc.add_heading(
    'Managing Sales, Purchases, Internal Transfers and Reporting '
    'Across Multiple Warehouses and Companies in Odoo 19', level=1)
title.runs[0].font.size = Pt(16)
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

meta = doc.add_paragraph('Rohan Raj | Feb 23, 2026')
meta.runs[0].font.size = Pt(10)
meta.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
meta.runs[0].font.italic = True

doc.add_paragraph(
    'This document explains how to manage daily operations in Odoo 19 when the business has '
    'multiple companies, branches and warehouses. It covers creating sales orders from specific '
    'warehouses, receiving purchases at branch locations, transferring stock between warehouses '
    '(manually and automatically), linking two companies so that a sales order in one company '
    'automatically creates a purchase order in the other, and viewing separate financial reports '
    'for each company and branch. The document is written for any business using Odoo 19 with '
    'this type of multi-company or multi-warehouse setup.'
)

# ══════════════════════════════════════════════════════════════════════
# UNDERSTANDING
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('Understanding the Setup', level=2)

doc.add_paragraph(
    'In Odoo 19, businesses can be structured in different ways depending on their legal and '
    'operational needs. The three main building blocks are companies, branches and warehouses.'
)

add_bold_para('Company', ' — A separate legal entity with its own GSTIN (or Tax ID), chart of accounts, '
    'bank accounts and financial statements. Each company files its own tax returns and maintains '
    'its own books. A business group with two different GSTINs needs two separate companies in Odoo.')

add_bold_para('Branch (Child Company)', ' — A subdivision within a parent company that shares the same '
    'GSTIN and chart of accounts. Branches are useful for tracking performance of different offices '
    'or divisions without maintaining fully separate books. In Odoo 19, child companies under a '
    'parent company serve as branches.')

add_bold_para('Warehouse', ' — A physical location where inventory is stored. Each warehouse belongs to '
    'exactly one company. A company can have multiple warehouses (for different stores, factories or '
    'storage facilities). Stock is tracked separately at each warehouse.')

add_image('screenshots/01_company_selector.png',
    'Company selector in the top right corner showing multiple companies and branches')

doc.add_paragraph(
    'The company selector in the top right corner of the screen allows switching between companies. '
    'When a company is selected, all operations, reports and data are filtered to that company only.'
)

# ══════════════════════════════════════════════════════════════════════
# PART 1: SALES
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('Part 1 — Creating a Sales Order from a Specific Warehouse', level=2)

doc.add_paragraph(
    'When a customer places an order, the business needs to select which warehouse will fulfil '
    'the delivery. This ensures goods are shipped from the correct location and stock levels '
    'are updated at the right warehouse.'
)

doc.add_paragraph(
    'To create a sales order, open the Sales application and click New. Select the customer, '
    'then add the products and quantities required. Before confirming the order, look for the '
    'Warehouse field on the sales order form. This field determines which warehouse the goods '
    'ship from. Click the Warehouse field and select the appropriate warehouse from the dropdown list.'
)

add_nav_path('Sales App → New → Select Customer → Add Products → Change Warehouse → Confirm')

add_image('screenshots/02_sales_order_warehouse.png',
    'Sales order form with the Warehouse field highlighted showing the dropdown list of available warehouses')

doc.add_paragraph(
    'Once the order is confirmed, Odoo automatically creates a delivery order from the selected '
    'warehouse. The invoice is generated under the company that owns the warehouse.'
)

add_tip(
    'To track which branch made the sale, use the Sales Team field on the sales order. '
    'Create a separate sales team for each branch so that sales can be filtered by branch in reports.'
)

# ══════════════════════════════════════════════════════════════════════
# PART 2: PURCHASES
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('Part 2 — Creating a Purchase Order for a Specific Warehouse', level=2)

doc.add_paragraph(
    'When goods are purchased from a vendor, they can be received directly into any warehouse. '
    'This avoids the need to receive everything at headquarters and then transfer.'
)

doc.add_paragraph(
    'To create a purchase order, open the Purchase application and click New. Select the vendor, '
    'then add the products and quantities. On the purchase order form, locate the Deliver To field. '
    'This determines which warehouse receives the incoming goods. Click on this field and select '
    'the destination warehouse.'
)

add_nav_path('Purchase App → New → Select Vendor → Add Products → Change Deliver To → Confirm Order')

add_image('screenshots/03_purchase_order_deliver_to.png',
    'Purchase order form with the Deliver To field highlighted showing the destination warehouse selection')

doc.add_paragraph(
    'After the order is confirmed, a receipt is created at the selected warehouse. When the goods '
    'arrive, open the receipt and click Validate to confirm that the products have been received.'
)

add_tip(
    'Make sure the correct company is selected in the company selector before creating the purchase '
    'order. The purchase bill will be recorded under the currently selected company.'
)

# ══════════════════════════════════════════════════════════════════════
# PART 3: INTERNAL TRANSFERS
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('Part 3 — Transferring Stock Between Warehouses (Same Company)', level=2)

doc.add_paragraph(
    'When stock needs to move from one warehouse to another within the same company, an internal '
    'transfer is used. Since both warehouses belong to the same company, this is a simple stock '
    'movement with no tax, no invoicing and no accounting entries.'
)

doc.add_heading('Method 1 — Manual Internal Transfer', level=3)

doc.add_paragraph(
    'To create a manual transfer, navigate to the Inventory application. Go to Operations and '
    'select Transfers. Click New to create a new transfer. Set the Operation Type to Internal '
    'Transfer, select the Source Location and the Destination Location, add the products and '
    'quantities, then click Validate when the goods are physically moved.'
)

add_nav_path('Inventory App → Operations → Transfers → New')

add_image('screenshots/04_internal_transfer.png',
    'Internal transfer form showing source location, destination location and product lines')

doc.add_heading('Method 2 — Automatic Resupply Between Warehouses', level=3)

doc.add_paragraph(
    'Instead of creating manual transfers each time, Odoo can be configured to automatically move '
    'stock from a supply warehouse to a branch warehouse whenever stock runs low. This requires '
    'three steps: enabling multi-step routes, configuring resupply on the branch warehouse, and '
    'creating reordering rules.'
)

add_bold_para('Step 1 — Enable Multi-Step Routes')

add_nav_path('Inventory App → Configuration → Settings → Warehouse section → Enable Multi-Step Routes → Save')

add_image('screenshots/05_multistep_routes.png',
    'Inventory settings showing the Multi-Step Routes checkbox enabled in the Warehouse section')

add_bold_para('Step 2 — Configure Resupply on the Branch Warehouse')

doc.add_paragraph(
    'Open the branch warehouse from the warehouse configuration screen. In the Resupply From field, '
    'check the box next to the warehouse that should supply this branch. Save the changes.'
)

add_nav_path('Inventory App → Configuration → Warehouses → Select the branch warehouse')

add_image('screenshots/06_warehouse_resupply.png',
    'Warehouse configuration showing the Resupply From field with the supply warehouse selected')

doc.add_paragraph(
    'Repeat this for any other warehouse that should be automatically resupplied.'
)

add_bold_para('Step 3 — Create Reordering Rules')

doc.add_paragraph(
    'Reordering rules tell Odoo when to trigger the automatic transfer. For each product at each '
    'branch warehouse, set the minimum quantity (trigger level) and maximum quantity (target level).'
)

add_nav_path('Inventory App → Operations → Replenishment → New')

add_image('screenshots/07_reordering_rules.png',
    'Reordering rules showing product, location, min/max quantities and the resupply route')

doc.add_paragraph(
    'When the scheduler runs (daily by default, or triggered manually from Inventory → Operations → '
    'Run Scheduler), Odoo checks all reordering rules and creates automatic transfers as needed.'
)

# ══════════════════════════════════════════════════════════════════════
# PART 4: INTER-COMPANY
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('Part 4 — Transferring Stock Between Two Companies (Inter-Company)', level=2)

doc.add_paragraph(
    'When two separate companies exist in the same Odoo database (for example with different '
    'GSTINs or tax registrations), stock cannot be moved using a simple internal transfer. Any '
    'goods movement between them is legally a sale from one company and a purchase by the other, '
    'requiring proper invoicing and tax.'
)

doc.add_paragraph(
    'Odoo provides a built-in feature to automate this. When a sales order is confirmed in '
    'Company A with Company B as the customer, Odoo can automatically create a corresponding '
    'purchase order in Company B with Company A as the vendor.'
)

doc.add_heading('Setting Up Automatic Inter-Company Transactions', level=3)

add_bold_para('Step 1 — Enable Inter-Company Transactions')

doc.add_paragraph(
    'Open the General Settings and enable the inter-company feature.'
)

add_nav_path('Settings → General Settings → Inter-Company Transactions → Enable → Save')

add_image('screenshots/08_intercompany_settings.png',
    'General Settings showing Inter-Company Transactions with auto-generation options enabled')

doc.add_paragraph(
    'Odoo installs the required modules automatically. The page reloads after approximately '
    '30 to 60 seconds.'
)

add_bold_para('Step 2 — Configure Rules for Each Company')

doc.add_paragraph(
    'After the module is installed, configure the inter-company rules for each company. Switch '
    'to each company using the company selector and then open the settings.'
)

add_nav_path('Company Selector → Select Company → Settings → General Settings → Inter-Company Transactions')

doc.add_paragraph('The following options can be configured for each company:')

# Options table
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Shading Accent 1'
table.rows[0].cells[0].text = 'Option'
table.rows[0].cells[1].text = 'What It Does'
for p in table.rows[0].cells[0].paragraphs:
    for r in p.runs: r.bold = True
for p in table.rows[0].cells[1].paragraphs:
    for r in p.runs: r.bold = True

opts = [
    ('Generate Purchase Orders', 'When this company confirms a sales order to another company, a purchase order is automatically created in the other company'),
    ('Generate Sales Orders', 'When this company confirms a purchase order from another company, a sales order is automatically created in the other company'),
    ('Generate Bills and Refunds', 'When this company posts an invoice for another company, a vendor bill is auto-created in the other company'),
    ('Create and Validate', 'The auto-generated document is also confirmed automatically instead of being left in draft'),
]
for i, (opt, desc) in enumerate(opts, 1):
    table.rows[i].cells[0].text = opt
    table.rows[i].cells[1].text = desc

doc.add_paragraph('')  # spacing

add_bold_para('Step 3 — Set the Warehouse for Auto-Generated Orders')

doc.add_paragraph(
    'In the inter-company settings, set the Use Warehouse field for each company. This determines '
    'which warehouse is used when Odoo auto-generates a purchase order or sales order in that company.'
)

doc.add_heading('How the Automatic Flow Works', level=3)

doc.add_paragraph(
    'After the configuration is complete, the inter-company flow works as follows:'
)

steps = [
    'In Company A, create a sales order with Company B as the customer. Add the products, quantities and prices. Confirm the sales order.',
    'Odoo automatically creates a purchase order in Company B with Company A as the vendor, matching the same products, quantities and prices.',
    'In Company A, process the delivery order (ship the goods). Click Validate.',
    'In Company B, process the receipt (receive the goods). Click Validate.',
    'In Company A, create and send the sales invoice with applicable tax.',
    'In Company B, record the vendor bill matching the purchase order.',
]
for step in steps:
    doc.add_paragraph(step, style='List Number')

add_image('screenshots/09_intercompany_flow.png',
    'Sales order confirmed in Company A automatically creates a matching purchase order in Company B')

add_tip(
    'Products must be shared between companies for this to work. On the product form, the Company '
    'field should be left blank (empty) so that both companies can access the same products.'
)

doc.add_heading('Manual Method (Without the Inter-Company Module)', level=3)

doc.add_paragraph(
    'If the inter-company module is not enabled, the same result can be achieved manually. Create '
    'a sales order in Company A with Company B as the customer. Then switch to Company B and create '
    'a purchase order with Company A as the vendor, adding the same products and quantities. Process '
    'the delivery, receipt and invoices separately in each company.'
)

# ══════════════════════════════════════════════════════════════════════
# PART 5: REPORTING
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('Part 5 — Viewing Financial Reports by Company', level=2)

doc.add_paragraph(
    'Odoo provides separate financial reports for each company. By switching companies using the '
    'company selector, the reports automatically show data for the selected company only.'
)

doc.add_heading('Profit and Loss Report', level=3)

doc.add_paragraph(
    'The Profit and Loss report shows whether a company is profitable during a specific period. '
    'To view it, first select the company using the company selector, then open the report.'
)

add_nav_path('Company Selector → Select the Company → Accounting App → Reporting → Profit and Loss')

add_image('screenshots/10_profit_loss.png',
    'Profit and Loss report showing revenue, cost of revenue, gross profit, operating expenses and net profit')

doc.add_paragraph(
    'The report displays revenue earned during the period, followed by the cost of revenue to '
    'calculate gross profit. Operating expenses are subtracted to show operating income or loss. '
    'Other income and expenses are then applied to show the net profit. To change the period, '
    'click the date button at the top of the report and select the required month, quarter, '
    'year or custom date range.'
)

doc.add_paragraph(
    'To view the report for a different company, switch to that company using the company selector '
    'and open the same report. Each company has its own separate financial data.'
)

doc.add_heading('Trial Balance with Account Grouping', level=3)

doc.add_paragraph(
    'The Trial Balance shows all account balances in a structured view. To see hierarchical grouping, '
    'click on Posted Entries within the report, enable Hierarchy and Subtotals and select Unfold All.'
)

add_nav_path('Accounting App → Reporting → Trial Balance → Options → Hierarchy and Subtotals → Unfold All')

add_image('screenshots/11_trial_balance.png',
    'Trial Balance with hierarchy enabled showing account groups expanding into sub-accounts')

doc.add_heading('Reports for Branches Using Analytic Accounts', level=3)

doc.add_paragraph(
    'Branches (child companies) that share the same GSTIN and chart of accounts as their parent '
    'do not have separate financial statements by default. All transactions run through the parent '
    'company. To track branch-level performance, analytic accounting is used.'
)

add_bold_para('Step 1 — Enable Analytic Accounting')

add_nav_path('Accounting App → Configuration → Settings → Analytics → Enable Analytic Accounting → Save')

add_image('screenshots/16_analytic_settings.png',
    'Accounting settings showing the Analytic Accounting checkbox enabled')

add_bold_para('Step 2 — Create Analytic Plan and Accounts for Branches')

add_nav_path('Accounting App → Configuration → Analytic Plans → New → Name: Branches')

doc.add_paragraph(
    'Within this plan, create an analytic account for each branch or warehouse location. When '
    'creating sales orders, purchase orders or journal entries, select the appropriate analytic '
    'account to tag the transaction to the correct branch.'
)

add_image('screenshots/12_analytic_plans.png',
    'Analytic plan called Branches with accounts for each branch and warehouse location')

add_bold_para('Step 3 — View Branch Profitability')

doc.add_paragraph(
    'Open the Profit and Loss report and filter by analytic account to see how much revenue '
    'and expense each branch has generated.'
)

add_nav_path('Accounting App → Reporting → Profit and Loss → Options → Filter by Analytic Account')

add_image('screenshots/17_pl_by_analytic.png',
    'Profit and Loss report broken down by branch showing revenue, costs and net profit per branch')

# ══════════════════════════════════════════════════════════════════════
# PART 6: INVENTORY REPORTS
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('Part 6 — Viewing Inventory Reports by Warehouse', level=2)

doc.add_heading('Current Stock by Warehouse', level=3)

doc.add_paragraph(
    'To see the stock available at each warehouse, open the inventory report and group by warehouse.'
)

add_nav_path('Inventory App → Reporting → Inventory Report → Group By → Warehouse')

add_image('screenshots/13_inventory_report.png',
    'Inventory report grouped by warehouse showing on-hand and forecasted quantities at each location')

doc.add_paragraph(
    'To check a specific product across all warehouses, open the product form and click the '
    'On Hand smart button.'
)

doc.add_heading('Stock Moves History', level=3)

add_nav_path('Inventory App → Reporting → Stock Moves')

doc.add_paragraph(
    'Filter by source location or destination location to see transfers involving a specific '
    'warehouse. For example, filter by source location to see all items sent out from a warehouse.'
)

# ══════════════════════════════════════════════════════════════════════
# PART 7: FUTURE GROWTH
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('Part 7 — Adding More Companies, Branches or Warehouses in the Future', level=2)

doc.add_paragraph(
    'As the business grows, new companies, branches or warehouses may need to be added. '
    'This section explains when and how to add each one.'
)

doc.add_heading('When to Add a New Company', level=3)

doc.add_paragraph(
    'Add a new company when the business creates a new legal entity with its own tax registration '
    '(different GSTIN or Tax ID), needs separate financial statements, or is required to file '
    'separate tax returns.'
)

add_nav_path('Settings → Users & Companies → Companies → New')

add_image('screenshots/14_new_company.png',
    'New company creation form with fields for company name, parent company, GSTIN and address')

doc.add_paragraph(
    'After creating the company, configure its chart of accounts and journals. If it needs to '
    'exchange goods with existing companies, enable inter-company transactions as described in Part 4.'
)

doc.add_heading('When to Add a New Branch', level=3)

doc.add_paragraph(
    'Add a new branch (child company) when a new office or department is created under an existing '
    'legal entity. The branch shares the same GSTIN and chart of accounts. Create it as a new '
    'company but set the Parent Company field to the existing parent. After adding, create an '
    'analytic account for the branch so transactions can be tracked.'
)

doc.add_heading('When to Add a New Warehouse', level=3)

doc.add_paragraph(
    'Add a new warehouse when a new physical store, factory or storage location is opened. '
    'Make sure the correct company is selected first.'
)

add_nav_path('Company Selector → Select Company → Inventory App → Configuration → Warehouses → New')

add_image('screenshots/15_new_warehouse.png',
    'New warehouse creation form with fields for name, short code, company and configuration')

doc.add_paragraph(
    'If the new warehouse should be automatically resupplied from an existing warehouse, enable '
    'the Resupply From field and create reordering rules as described in Part 3.'
)

# Decision table
doc.add_heading('Quick Decision Guide', level=3)

table2 = doc.add_table(rows=7, cols=3)
table2.style = 'Light Shading Accent 1'
for i, h in enumerate(['Situation', 'What to Add', 'Key Action']):
    table2.rows[0].cells[i].text = h
    for p in table2.rows[0].cells[i].paragraphs:
        for r in p.runs: r.bold = True

decisions = [
    ('New store, same legal entity', 'New Warehouse', 'Create warehouse under existing company'),
    ('New legal entity, different GSTIN', 'New Company', 'Create company with own chart of accounts'),
    ('New department, same entity', 'New Branch', 'Create child company under parent'),
    ('New store needs auto-stock', 'Warehouse + Resupply', 'Create warehouse, enable Resupply From'),
    ('New company trades with existing', 'Company + Inter-Company', 'Enable inter-company transactions'),
    ('Track profitability per location', 'Analytic Account', 'Add analytic account, tag transactions'),
]
for i, (sit, what, action) in enumerate(decisions, 1):
    table2.rows[i].cells[0].text = sit
    table2.rows[i].cells[1].text = what
    table2.rows[i].cells[2].text = action

# ══════════════════════════════════════════════════════════════════════
# PART 8: QUICK REFERENCE
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('Part 8 — Quick Reference Summary', level=2)

table3 = doc.add_table(rows=11, cols=2)
table3.style = 'Light Shading Accent 1'
table3.rows[0].cells[0].text = 'Operation'
table3.rows[0].cells[1].text = 'Navigation Path'
for p in table3.rows[0].cells[0].paragraphs:
    for r in p.runs: r.bold = True
for p in table3.rows[0].cells[1].paragraphs:
    for r in p.runs: r.bold = True

ref_data = [
    ('Sale from a specific warehouse', 'Sales → New → Set Warehouse → Confirm'),
    ('Purchase for a specific warehouse', 'Purchase → New → Set Deliver To → Confirm'),
    ('Manual internal transfer', 'Inventory → Operations → Transfers → New'),
    ('Enable automatic resupply', 'Inventory → Configuration → Warehouses → Resupply From'),
    ('Set up reordering rules', 'Inventory → Operations → Replenishment → New'),
    ('Enable inter-company auto SO/PO', 'Settings → General Settings → Inter-Company Transactions'),
    ('Profit and Loss by company', 'Switch company → Accounting → Reporting → Profit and Loss'),
    ('Branch profitability', 'Accounting → Reporting → P&L → Filter by Analytic Account'),
    ('Stock by warehouse', 'Inventory → Reporting → Inventory Report → Group By Warehouse'),
    ('Add a new warehouse', 'Inventory → Configuration → Warehouses → New'),
]
for i, (op, path) in enumerate(ref_data, 1):
    table3.rows[i].cells[0].text = op
    table3.rows[i].cells[1].text = path

doc.add_paragraph('')
doc.add_paragraph(
    'This document covers the essential daily operations and configuration steps for managing '
    'multiple companies, branches and warehouses in Odoo 19. As the business grows, the same '
    'principles apply — add warehouses for new physical locations, branches for departments within '
    'the same legal entity, and companies for new legal entities. For additional configuration '
    'such as user access restrictions per warehouse, E-Waybill generation for inter-state transfers, '
    'or manufacturing integration, contact the implementation team.'
)

# Save
doc.save('CLIENT_OPERATIONS_GUIDE_v2.docx')
print("Document saved: CLIENT_OPERATIONS_GUIDE_v2.docx")
