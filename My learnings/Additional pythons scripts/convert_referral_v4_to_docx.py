#!/usr/bin/env python3
"""
Convert BLOG_Referral_App_Odoo_19_V4_REPORTS_CORRECTED.md to a professional Word (.docx) document.
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT


def setup_styles(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15

    heading_configs = {
        'Heading 1': (Pt(22), RGBColor(0x1A, 0x1A, 0x2E), True, Pt(18), Pt(8)),
        'Heading 2': (Pt(17), RGBColor(0x2C, 0x3E, 0x50), True, Pt(14), Pt(6)),
        'Heading 3': (Pt(14), RGBColor(0x34, 0x49, 0x5E), True, Pt(10), Pt(4)),
        'Heading 4': (Pt(12), RGBColor(0x5D, 0x6D, 0x7E), True, Pt(8), Pt(4)),
    }
    for name, (size, color, bold, sp_before, sp_after) in heading_configs.items():
        if name in doc.styles:
            s = doc.styles[name]
            s.font.name = 'Calibri'
            s.font.size = size
            s.font.color.rgb = color
            s.font.bold = bold
            s.paragraph_format.space_before = sp_before
            s.paragraph_format.space_after = sp_after


def add_formatted_text(paragraph, text):
    combined = r'(\*\*\*(.*?)\*\*\*|\*\*(.*?)\*\*|__(.*?)__|`(.*?)`|\[(.*?)\]\((.*?)\)|\*(.*?)\*)'
    pos = 0
    for m in re.finditer(combined, text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        matched = m.group()
        if matched.startswith('***'):
            run = paragraph.add_run(m.group(2))
            run.bold = True; run.italic = True
        elif matched.startswith('**'):
            run = paragraph.add_run(m.group(3))
            run.bold = True
        elif matched.startswith('__'):
            run = paragraph.add_run(m.group(4))
            run.bold = True
        elif matched.startswith('`'):
            run = paragraph.add_run(m.group(5))
            run.font.name = 'Consolas'; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif matched.startswith('['):
            run = paragraph.add_run(m.group(6))
            run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC); run.underline = True
        elif matched.startswith('*'):
            run = paragraph.add_run(m.group(8))
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def convert_md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    setup_styles(doc)

    i = 0
    in_table = False
    table_rows = []

    def flush_table():
        nonlocal table_rows, in_table
        if not table_rows:
            in_table = False
            return
        parsed = []
        for row in table_rows:
            cells = [c.strip() for c in row.strip('|').split('|')]
            if all(set(c.strip()) <= set('-: ') for c in cells):
                continue
            parsed.append(cells)
        if not parsed:
            in_table = False; table_rows = []; return
        max_cols = max(len(r) for r in parsed)
        for r in parsed:
            while len(r) < max_cols: r.append('')
        table = doc.add_table(rows=len(parsed), cols=max_cols)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for ri, row_data in enumerate(parsed):
            for ci, cell_text in enumerate(row_data):
                cell = table.cell(ri, ci)
                cell.text = ''
                p = cell.paragraphs[0]
                add_formatted_text(p, cell_text)
                for run in p.runs: run.font.size = Pt(10)
                if ri == 0:
                    for run in p.runs: run.bold = True
        doc.add_paragraph('')
        table_rows = []; in_table = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            if in_table: flush_table()
            i += 1; continue
        if stripped.startswith('[Image:') or stripped.startswith('!['):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.size = Pt(9); run.font.color.rgb = RGBColor(0x99, 0x99, 0x99); run.italic = True
            i += 1; continue
        if stripped.startswith('|') and stripped.endswith('|'):
            in_table = True; table_rows.append(stripped); i += 1; continue
        elif in_table:
            flush_table()
        if stripped.startswith('#'):
            match = re.match(r'^(#{1,4})\s+(.*)', stripped)
            if match:
                level = len(match.group(1)); heading_text = match.group(2)
                p = doc.add_heading(level=level)
                add_formatted_text(p, heading_text)
                i += 1; continue
        num_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if num_match:
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, num_match.group(2))
            i += 1; continue
        if stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, stripped[2:])
            i += 1; continue
        if '=>' in stripped and not stripped.startswith('#'):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(stripped)
            run.font.name = 'Consolas'; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)
            i += 1; continue
        p = doc.add_paragraph()
        add_formatted_text(p, stripped)
        i += 1

    if in_table: flush_table()
    doc.save(docx_path)
    print(f"Successfully created: {docx_path}")


if __name__ == "__main__":
    base = r"c:\Odoo Study"
    md_path = os.path.join(base, "BLOG_Referral_App_Odoo_19_V4_REPORTS_CORRECTED.md")
    docx_path = os.path.join(base, "My learnings", "Rohan_Documentation", "BLOG_Referral_App_V4_FINAL.docx")
    convert_md_to_docx(md_path, docx_path)
