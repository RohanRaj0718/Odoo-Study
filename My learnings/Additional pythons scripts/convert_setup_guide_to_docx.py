"""Convert PSI_COMPLETE_SETUP_GUIDE.md to a formatted .docx file."""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

INPUT_FILE = "PSI_COMPLETE_SETUP_GUIDE.md"
OUTPUT_FILE = "PSI_COMPLETE_SETUP_GUIDE.docx"


def set_cell_shading(cell, color_hex):
    """Set background shading for a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_paragraph(doc, text, style_name, bold=False, color=None, size=None, space_after=None):
    """Add a paragraph with optional formatting."""
    p = doc.add_paragraph()
    p.style = doc.styles[style_name] if style_name in [s.name for s in doc.styles] else doc.styles['Normal']
    run = p.add_run(text)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def parse_table(lines):
    """Parse markdown table lines into headers and rows."""
    headers = []
    rows = []
    for i, line in enumerate(lines):
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        if i == 0:
            headers = cells
        elif i == 1:
            continue  # separator line
        else:
            rows.append(cells)
    return headers, rows


def add_table_to_doc(doc, headers, rows):
    """Add a formatted table to the document."""
    if not headers:
        return
    num_cols = len(headers)
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "1F4E79")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            if i < num_cols:
                cell = row.cells[i]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    doc.add_paragraph()  # spacing after table


def process_markdown(doc, md_text):
    """Process the markdown and build the docx."""
    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # Code block start/end
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block — flush
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                # Light gray background via shading
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0"/>')
                p._p.get_or_add_pPr().append(shading)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped == '---':
            doc.add_paragraph('─' * 70)
            i += 1
            continue

        # Headings
        if stripped.startswith('# ') and not stripped.startswith('## '):
            text = stripped[2:].strip()
            p = doc.add_heading(text, level=0)
            i += 1
            continue

        if stripped.startswith('## '):
            text = stripped[3:].strip()
            doc.add_heading(text, level=1)
            i += 1
            continue

        if stripped.startswith('### '):
            text = stripped[4:].strip()
            doc.add_heading(text, level=2)
            i += 1
            continue

        if stripped.startswith('#### '):
            text = stripped[5:].strip()
            doc.add_heading(text, level=3)
            i += 1
            continue

        # Table detection
        if '|' in stripped and stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 2:
                headers, rows = parse_table(table_lines)
                add_table_to_doc(doc, headers, rows)
            continue

        # Blockquote
        if stripped.startswith('> '):
            text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            # Handle bold within blockquote
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                else:
                    run = p.add_run(part)
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x66, 0x00, 0x00)
            i += 1
            continue

        # Numbered list items (from TOC or instructions)
        num_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if num_match:
            text = num_match.group(2)
            # Remove markdown links [text](link)
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            p = doc.add_paragraph(text, style='List Number')
            for run in p.runs:
                run.font.size = Pt(10)
            i += 1
            continue

        # Bullet points
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            p = doc.add_paragraph(style='List Bullet')
            # Handle bold/code in bullet
            parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.size = Pt(10)
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
                else:
                    run = p.add_run(part)
                    run.font.size = Pt(10)
            i += 1
            continue

        # Regular paragraph — handle inline formatting
        p = doc.add_paragraph()
        text = stripped
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
                run.font.size = Pt(10)
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
            else:
                run = p.add_run(part)
                run.font.size = Pt(10)

        i += 1


def main():
    # Read markdown
    with open(INPUT_FILE, 'r', encoding='utf-8') as f:
        md_text = f.read()

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Process the markdown
    process_markdown(doc, md_text)

    # Save
    doc.save(OUTPUT_FILE)
    print(f"✅ Created: {OUTPUT_FILE}")
    print(f"   Pages will be approximately 25-30 in Word")


if __name__ == '__main__':
    main()
