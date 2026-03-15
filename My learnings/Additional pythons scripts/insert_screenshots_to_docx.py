#!/usr/bin/env python3
"""
Insert captured screenshots into the corrected Subcontracting blog DOCX.
Replaces all [Image: ...] placeholder paragraphs with actual images.
"""
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX_PATH = r'C:\Odoo Study\BLOG_Subcontracting_PUBLISH_READY_CORRECTED.docx'
OUTPUT_PATH = r'C:\Odoo Study\BLOG_Subcontracting_PUBLISH_READY_CORRECTED.docx'
SCREENSHOT_DIR = r'C:\Odoo Study\blog_screenshots'

# Mapping: placeholder text -> screenshot filename
SCREENSHOT_MAP = {
    'Subcontracting Settings': '01_subcontracting_settings.png',
    'Subcontractor Contact Form': '02_subcontractor_contact.png',
    'Product Form': '03_product_form.png',
    'BoM Type Set to Subcontracting': '04_bom_subcontracting_type.png',
    'Components Tab': '05_bom_components.png',
    'Vendor Pricelist on Purchase Tab': '06_vendor_pricelist.png',
    'Purchase Order for Subcontracted Product': '07_purchase_order.png',
    'Smart Buttons on Purchase Order': '08_smart_buttons.png',
    'Resupply Transfer with Components': '09_resupply_transfer.png',
    'Subcontracting Receipt': '10_subcontracting_receipt.png',
    'Dropshipping Settings': '11_dropshipping_settings.png',
    'Dropship Subcontracting Flow': '12_dropship_flow.png',
}

doc = Document(DOCX_PATH)

replaced = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # Match [Image: XYZ] format (from docx converter)
    if text.startswith('[Image:') and text.endswith(']'):
        label = text[7:-1].strip()  # Extract "XYZ" from "[Image: XYZ]"
        
        if label in SCREENSHOT_MAP:
            img_path = os.path.join(SCREENSHOT_DIR, SCREENSHOT_MAP[label])
            if os.path.exists(img_path):
                # Clear the placeholder text
                para.clear()
                
                # Add the image
                run = para.add_run()
                run.add_picture(img_path, width=Inches(6.0))
                
                # Center the image
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                replaced += 1
                print(f'  Inserted: {label} -> {SCREENSHOT_MAP[label]}')
            else:
                print(f'  MISSING: {img_path}')
        else:
            print(f'  NO MAPPING for: "{label}"')
    
    # Also check for Infintor banner image (keep as-is)

doc.save(OUTPUT_PATH)
size = os.path.getsize(OUTPUT_PATH)
print(f'\nDone! Replaced {replaced} placeholders with screenshots.')
print(f'Output: {OUTPUT_PATH} ({size:,} bytes)')
