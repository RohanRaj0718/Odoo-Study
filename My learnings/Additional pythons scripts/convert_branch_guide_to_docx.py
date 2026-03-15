"""Convert PSI_BRANCH_WAREHOUSE_OPERATIONS_GUIDE.md to a formatted .docx file."""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

INPUT_FILE = "PSI_BRANCH_WAREHOUSE_OPERATIONS_GUIDE.md"
OUTPUT_FILE = "PSI_BRANCH_WAREHOUSE_OPERATIONS_GUIDE.docx"


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_runs(paragraph, text):
    """Parse inline markdown formatting and add runs."""
    # Split by bold markers **text** and `code`
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            paragraph.add_run(part)


def parse_table(lines):
    headers = []
    rows = []
    for i, line in enumerate(lines):
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        if i == 0:
            headers = cells
        elif i == 1:
            continue
        else:
            rows.append(cells)
    return headers, rows


def convert_md_to_docx(input_file, output_file):
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # Skip horizontal rules
        if line.strip() == '---':
            doc.add_paragraph()  # spacer
            i += 1
            continue

        # Code block start/end
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block - add collected lines
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                # Add light gray background via shading
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
                p._element.get_or_add_pPr().append(shading)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Headings
        if line.startswith('# ') and not line.startswith('## '):
            p = doc.add_heading(line[2:].strip(), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue

        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        # Blockquote
        if line.strip().startswith('> '):
            text = line.strip()[2:]
            # Collect multi-line blockquotes
            while i + 1 < len(lines) and lines[i + 1].strip().startswith('> '):
                i += 1
                text += ' ' + lines[i].strip()[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(text.replace('**', ''))
            run.bold = True
            run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
            run.font.size = Pt(11)
            i += 1
            continue

        # Table detection
        if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            table_lines = [line]
            j = i + 1
            while j < len(lines) and '|' in lines[j] and lines[j].strip():
                table_lines.append(lines[j])
                j += 1

            headers, rows = parse_table(table_lines)
            if headers and rows:
                num_cols = len(headers)
                table = doc.add_table(rows=1 + len(rows), cols=num_cols)
                table.style = 'Light Grid Accent 1'

                # Headers
                for ci, h in enumerate(headers):
                    cell = table.rows[0].cells[ci]
                    cell.text = h.replace('**', '')
                    set_cell_shading(cell, "1F4E79")
                    for par in cell.paragraphs:
                        for run in par.runs:
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.bold = True
                            run.font.size = Pt(10)

                # Data rows
                for ri, row_data in enumerate(rows):
                    for ci in range(min(len(row_data), num_cols)):
                        cell = table.rows[ri + 1].cells[ci]
                        cell.text = row_data[ci].replace('**', '')
                        for par in cell.paragraphs:
                            for run in par.runs:
                                run.font.size = Pt(10)

                doc.add_paragraph()  # spacer after table

            i = j
            continue

        # Bullet points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_runs(p, text)
            i += 1
            continue

        # Numbered list
        num_match = re.match(r'^(\d+)\.\s+(.*)', line.strip())
        if num_match:
            text = num_match.group(2)
            p = doc.add_paragraph(style='List Number')
            add_formatted_runs(p, text)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_runs(p, line.strip())
        i += 1

    doc.save(output_file)
    print(f"Successfully created: {output_file}")


if __name__ == '__main__':
    convert_md_to_docx(INPUT_FILE, OUTPUT_FILE)
