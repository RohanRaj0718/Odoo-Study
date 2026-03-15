import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def parse_md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Style headings
    for level in range(1, 5):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.color.rgb = RGBColor(0, 51, 102)
        heading_style.font.name = 'Calibri'
    
    doc.styles['Heading 1'].font.size = Pt(20)
    doc.styles['Heading 2'].font.size = Pt(16)
    doc.styles['Heading 3'].font.size = Pt(14)
    doc.styles['Heading 4'].font.size = Pt(12)

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []
    
    def flush_table(doc, table_rows):
        if not table_rows:
            return
        # Parse table rows
        parsed = []
        for row in table_rows:
            cells = [c.strip() for c in row.strip('|').split('|')]
            # skip separator rows
            if all(set(c.strip()) <= set('-: ') for c in cells):
                continue
            parsed.append(cells)
        
        if not parsed:
            return
            
        max_cols = max(len(r) for r in parsed)
        # Pad rows
        for r in parsed:
            while len(r) < max_cols:
                r.append('')
        
        table = doc.add_table(rows=len(parsed), cols=max_cols)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        for ri, row_data in enumerate(parsed):
            for ci, cell_text in enumerate(row_data):
                cell = table.cell(ri, ci)
                cell.text = ''
                p = cell.paragraphs[0]
                # Handle bold markers
                add_formatted_text(p, cell_text)
                p.style.font.size = Pt(10)
        
        doc.add_paragraph('')  # spacing after table
    
    def add_formatted_text(paragraph, text):
        """Add text with basic bold/italic/code formatting."""
        # Process bold, italic, code
        parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(128, 0, 0)
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                paragraph.add_run(part)
    
    def add_code_block(doc, code_lines):
        code_text = '\n'.join(code_lines)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 100, 0)
    
    while i < len(lines):
        line = lines[i]
        raw = line.rstrip('\n')
        stripped = raw.strip()
        
        # Code block toggle
        if stripped.startswith('```'):
            if in_code_block:
                # End code block
                add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                # Flush table if pending
                if in_table:
                    flush_table(doc, table_rows)
                    table_rows = []
                    in_table = False
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(raw)
            i += 1
            continue
        
        # Table detection
        if stripped.startswith('|') and '|' in stripped[1:]:
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(stripped)
            i += 1
            continue
        else:
            if in_table:
                flush_table(doc, table_rows)
                table_rows = []
                in_table = False
        
        # Empty line
        if not stripped:
            i += 1
            continue
        
        # Headings
        if stripped.startswith('#'):
            match = re.match(r'^(#{1,6})\s+(.*)', stripped)
            if match:
                level = min(len(match.group(1)), 4)
                heading_text = match.group(2).strip()
                # Remove markdown formatting from heading
                heading_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', heading_text)
                heading_text = re.sub(r'\*([^*]+)\*', r'\1', heading_text)
                heading_text = re.sub(r'`([^`]+)`', r'\1', heading_text)
                doc.add_heading(heading_text, level=level)
                i += 1
                continue
        
        # Blockquotes
        if stripped.startswith('>'):
            quote_text = stripped.lstrip('> ').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Check for multiline blockquote
            while i + 1 < len(lines) and lines[i + 1].strip().startswith('>'):
                i += 1
                next_line = lines[i].strip().lstrip('> ').strip()
                if next_line:
                    quote_text += '\n' + next_line
            
            add_formatted_text(p, quote_text)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(80, 80, 80)
            i += 1
            continue
        
        # Horizontal rule
        if stripped in ['---', '***', '___']:
            p = doc.add_paragraph('─' * 60)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.color.rgb = RGBColor(180, 180, 180)
            i += 1
            continue
        
        # Bullet points
        if re.match(r'^[-*]\s', stripped) or re.match(r'^\d+\.\s', stripped):
            # Determine list type
            if re.match(r'^\d+\.\s', stripped):
                text = re.sub(r'^\d+\.\s+', '', stripped)
                p = doc.add_paragraph(style='List Number')
            else:
                text = re.sub(r'^[-*]\s+', '', stripped)
                p = doc.add_paragraph(style='List Bullet')
            
            # Handle indented sub-items
            indent_level = len(raw) - len(raw.lstrip())
            if indent_level >= 4:
                p.paragraph_format.left_indent = Inches(0.5 + (indent_level // 2) * 0.25)
            
            p.clear()
            add_formatted_text(p, text)
            i += 1
            continue
        
        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_text(p, stripped)
        i += 1
    
    # Flush remaining
    if in_table:
        flush_table(doc, table_rows)
    if code_lines:
        add_code_block(doc, code_lines)
    
    doc.save(docx_path)
    print(f"Created: {docx_path}")


# Convert all files
base = r"c:\Odoo Study\MANUFACTURING"

files_to_convert = [
    ("BOM_TWO_KEY_TOPICS_GUIDE.md", "BOM_TWO_KEY_TOPICS_GUIDE.docx"),
    ("BA_COMPLETE_CLIENT_ANALYSIS.md", "BA_COMPLETE_CLIENT_ANALYSIS.docx"),
    ("HONEST_GAP_ANALYSIS.md", "HONEST_GAP_ANALYSIS.docx"),
    ("GAPS_QUICK_SUMMARY.md", "GAPS_QUICK_SUMMARY.docx")
]

for md_file, docx_file in files_to_convert:
    md_path = os.path.join(base, md_file)
    docx_path = os.path.join(base, docx_file)
    if os.path.exists(md_path):
        parse_md_to_docx(md_path, docx_path)
    else:
        print(f"Skipped: {md_file} (not found)")

print("\nDone! All Word files created.")
