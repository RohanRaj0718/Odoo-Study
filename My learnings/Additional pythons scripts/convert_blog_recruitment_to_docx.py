from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

md_path = 'BLOG_Recruitment_PUBLISH_READY.md'
docx_path = 'BLOG_Recruitment_PUBLISH_READY.docx'

with open(md_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

# Create document
doc = Document()

# Style helpers
def add_heading(text, level=1):
    doc.add_heading(text, level)

def add_paragraph(text, bold=False, italic=False, size=12, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if align:
        p.alignment = align
    return p

# Markdown to docx conversion
for line in lines:
    line = line.rstrip()
    if line.startswith('# '):
        add_heading(line[2:], level=1)
    elif line.startswith('##### '):
        add_heading(line[6:], level=2)
    elif line.startswith('!['):  # Image placeholder
        m = re.match(r'!\[(.*?)\]\((.*?)\)', line)
        if m:
            desc = m.group(1)
            add_paragraph(f'[Image: {desc}]', italic=True, size=11)
    elif line.startswith('['):  # Footer link
        m = re.match(r'\[(.*?)\]\((.*?)\)', line)
        if m:
            desc = m.group(1)
            url = m.group(2)
            add_paragraph(f'{desc}: {url}', italic=True, size=11)
    elif line.strip() == '':
        doc.add_paragraph('')
    else:
        add_paragraph(line)

# Save
doc.save(docx_path)
print(f'DOCX file created: {docx_path}')
