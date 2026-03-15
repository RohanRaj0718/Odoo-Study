"""
Edit Multiple_Branches.docx — Rewrite Part 4 (Inter-branch/Inter-company transfers)
Based on actual Odoo 19 database findings:
  - Branches can transfer stock via Inter-company transit location (NO SO/PO needed)
  - SO/PO method is only required for truly separate legal entities
"""

from docx import Document
from docx.shared import Pt, RGBColor
from copy import deepcopy
import os

doc = Document('Multiple_Branches.docx')

# ─── Identify the range to replace ───
# Part 4 starts at paragraph 63, Part 5 starts at paragraph 84
# We need to replace paragraphs 63-83 (inclusive)

START_IDX = 63  # "Part 4 — Transferring Stock Between Companies or Branches"
END_IDX = 83    # Last paragraph before Part 5 ("Manual Method..." paragraph)

# ─── Helper to create formatted paragraphs ───
def add_para(doc, text, style='Normal', bold=False, index=None):
    """Insert a paragraph at a specific index or append."""
    p = doc.add_paragraph(text, style=style)
    if bold:
        for run in p.runs:
            run.bold = True
    return p

def make_bold_normal(para_element, text):
    """Create a new paragraph element with bold text in Normal style."""
    from docx.oxml.ns import qn
    from lxml import etree
    p = deepcopy(para_element)
    # Clear existing runs
    for r in p.findall(qn('w:r')):
        p.remove(r)
    # Add new run with bold
    r = etree.SubElement(p, qn('w:r'))
    rpr = etree.SubElement(r, qn('w:rPr'))
    b = etree.SubElement(rpr, qn('w:b'))
    t = etree.SubElement(r, qn('w:t'))
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    return p

# ─── Build new content for Part 4 ───
# Strategy: Delete old paragraphs 64-83, then modify paragraph 63 and insert new ones after it

# First, let's get the XML body element
body = doc.element.body

# Collect the paragraph elements to remove (indices 64 through 83)
paras_to_remove = []
for i in range(64, END_IDX + 1):
    paras_to_remove.append(doc.paragraphs[i]._element)

# Remove them
for p_elem in paras_to_remove:
    body.remove(p_elem)

# Now paragraph 63 is the Part 4 heading - it stays
# We need to insert new paragraphs AFTER paragraph 63
# After removal, the paragraph after 63 is what was paragraph 84 (Part 5)

anchor = doc.paragraphs[63]._element  # Part 4 heading

# ─── New content to insert ───
new_content = [
    # Intro paragraph
    ("Normal", False, [
        (False, "When stock needs to move between different branches or companies, there are "),
        (True, "two methods"),
        (False, " available in Odoo 19. The method you choose depends on whether the entities are branches under the same parent company or completely separate legal entities with different GSTINs.")
    ]),
    
    # Blank line
    ("Normal", False, [("", False, "")]),
    
    # ─── Method 1 heading ───
    ("Heading 3", False, [(False, "Method 1 — Internal Transfer via Inter-Company Transit (For Branches)")]),
    
    # Explanation
    ("Normal", False, [
        (False, "When branches share the same parent company (same GSTIN/Tax ID), stock can be transferred using a simple "),
        (True, "internal transfer through the Inter-company transit location"),
        (False, ". This is a built-in location in Odoo that has no company assigned to it, making it accessible by all companies and branches. This method requires "),
        (True, "no sales order, no purchase order, no invoice and no tax"),
        (False, ". Products move at their existing cost price with zero financial impact.")
    ]),
    
    # How it works
    ("Normal", True, [(True, "How It Works")]),
    
    ("Normal", False, [
        (False, "The transfer happens in two steps. First, the sending branch creates an internal transfer from its warehouse stock location to the Inter-company transit location. Second, the receiving branch creates an internal transfer from the Inter-company transit location to its warehouse stock location.")
    ]),
    
    # Flow diagram text
    ("Normal", True, [(True, "Transfer Flow:")]),
    
    ("Normal", False, [
        (False, "Branch A Warehouse → Inter-company Transit → Branch B Warehouse")
    ]),
    
    # Step by step
    ("Normal", True, [(True, "Step 1 — Send from Source Branch")]),
    
    ("Normal", False, [
        (False, "Switch to the sending branch using the company selector. Open the Inventory application, go to Operations and select Internal Transfers. Click New. Set the Source Location to the branch warehouse stock location (for example WH/Stock). Set the Destination Location to Inter-company transit. Add the products and quantities. Click Validate.")
    ]),
    
    ("Normal", False, [
        (True, "Inventory App → Operations → Internal Transfers → New → Source: WH/Stock → Destination: Inter-company transit → Add Products → Validate")
    ]),
    
    # Step 2
    ("Normal", True, [(True, "Step 2 — Receive at Destination Branch")]),
    
    ("Normal", False, [
        (False, "Switch to the receiving branch using the company selector. Open the Inventory application, go to Operations and select Internal Transfers. Click New. Set the Source Location to Inter-company transit. Set the Destination Location to the receiving branch warehouse stock location (for example DF/Stock). Add the same products and quantities. Click Validate.")
    ]),
    
    ("Normal", False, [
        (True, "Inventory App → Operations → Internal Transfers → New → Source: Inter-company transit → Destination: DF/Stock → Add Products → Validate")
    ]),
    
    # Important notes
    ("Normal", True, [(True, "Key Points About This Method:")]),
    
    ("Normal", False, [
        (False, "No journal entries are created when products use periodic (manual) valuation, which is the default. Product sale price and cost price remain exactly the same after the transfer. The Inter-company transit location is automatically created by Odoo when multi-company is enabled. It has no company assigned (company field is empty), which is why both branches can access it. After the transfer, the transit location should show zero quantity as all goods have moved to the destination.")
    ]),
    
    # When to use
    ("Normal", True, [(True, "When to Use This Method:")]),
    
    ("Normal", False, [
        (False, "Use this method when transferring stock between branches that share the same parent company and the same GSTIN. Since branches are subdivisions of the same legal entity, there is no legal requirement to create sales and purchase documents between them. This is similar to moving goods between two shelves or two stores of the same business.")
    ]),
    
    # Blank
    ("Normal", False, [("", False, "")]),
    
    # ─── Method 2 heading ───
    ("Heading 3", False, [(False, "Method 2 — Inter-Company Transactions via SO/PO (For Separate Legal Entities)")]),
    
    ("Normal", False, [
        (False, "When the entities are "),
        (True, "separate legal entities with different GSTINs"),
        (False, ", the transfer must be treated as a sale from one company and a purchase by the other. This is required by law because each entity files its own tax returns and must maintain proper documentation for goods movement between them.")
    ]),
    
    # Setting up
    ("Normal", True, [(True, "Setting Up Automatic Inter-Company Transactions")]),
    
    ("Normal", False, [(True, "Step 1 — Enable Inter-Company Transactions")]),
    ("Normal", False, [(False, "Open the General Settings and enable the inter-company feature.")]),
    ("Normal", False, [(True, "Settings → General Settings → Inter-Company Transactions → Enable → Save")]),
    
    # Step 2
    ("Normal", False, [(True, "Step 2 — Configure Rules for Each Company")]),
    ("Normal", False, [
        (False, "Switch to each company using the company selector and configure the inter-company rules in the settings. The following options can be configured for each company:")
    ]),
    
    # Note about table - the existing table stays, we just reference it
    
    # Step 3
    ("Normal", False, [(True, "Step 3 — Set the Purchase Journal, Warehouse and Operation Type for Auto-Generated Orders")]),
    ("Normal", False, [
        (False, "In the inter-company settings, set the Purchase Journal, Use Warehouse and Use Operation field for each company. This determines which purchase journal, warehouse and operation type is used when Odoo auto-generates a purchase order or sales order in that company.")
    ]),
    
    # Workflow
    ("Normal", False, [
        (False, "When a sales order is confirmed in Company A with Company B as the customer, Odoo automatically creates a corresponding purchase order in Company B with Company A as the vendor. The delivery, receipt and invoicing follow the standard process in each company.")
    ]),
    
    # Manual method
    ("Heading 3", False, [(False, "Manual Method (Without the Inter-Company Module)")]),
    
    ("Normal", False, [
        (False, "If the inter-company module is not enabled, the same result can be achieved manually. Create a sales order in Company A with Company B as the customer. Then switch to Company B and create a purchase order with Company A as the vendor, adding the same products and quantities. Process the delivery, receipt and invoices separately in each company.")
    ]),
    
    # Comparison
    ("Heading 3", False, [(False, "Choosing the Right Method")]),

    ("Normal", False, [
        (False, "Use the Internal Transfer via Inter-company transit method when both entities are branches under the same parent company with the same GSTIN. This is simpler, faster and creates no unnecessary accounting entries. Use the Inter-Company SO/PO method only when the entities are separate legal entities with different GSTINs, where tax law requires proper sale and purchase documentation.")
    ]),
]

# ─── Insert new paragraphs after the anchor ───
from docx.oxml.ns import qn
from lxml import etree

# Get reference to the element AFTER the anchor (which is now Part 5 heading)
# We'll insert before that element
next_elem = anchor.getnext()

for style_name, all_bold, runs_data in new_content:
    # Create new paragraph element
    new_p = doc.add_paragraph()
    new_p.style = doc.styles[style_name]
    
    # Clear the auto-added empty run
    for r in new_p.runs:
        r.text = ""
    new_p.clear()
    
    # Add runs
    for item in runs_data:
        if len(item) == 3:
            is_bold, _, text = item
            if not text:
                continue
        else:
            is_bold, text = item
        
        run = new_p.add_run(text)
        if is_bold or all_bold:
            run.bold = True
    
    # Move the paragraph element from end of document to the correct position
    p_elem = new_p._element
    body.remove(p_elem)
    if next_elem is not None:
        next_elem.addprevious(p_elem)
    else:
        body.append(p_elem)

# ─── Save ───
output_path = 'Multiple_Branches.docx'
doc.save(output_path)
print(f"✅ Saved updated document: {output_path}")
print(f"   File size: {os.path.getsize(output_path) / 1024:.1f} KB")

# ─── Verify by reading back ───
print("\n=== VERIFICATION: Part 4 section in updated document ===")
doc2 = Document(output_path)
in_part4 = False
for i, p in enumerate(doc2.paragraphs):
    text = p.text.strip()
    style = p.style.name if p.style else ''
    if 'Part 4' in text:
        in_part4 = True
    if 'Part 5' in text:
        in_part4 = False
        break
    if in_part4 and text:
        print(f"  [{i:3d}] ({style}) {text[:120]}")
