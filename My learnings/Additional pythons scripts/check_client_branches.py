#!/usr/bin/env python3
"""
Read both branch-related docx files and check client-cient database setup.
"""

import xmlrpc.client
from docx import Document
import os

# ── Read DOCX files ──
def read_docx(path):
    doc = Document(path)
    text = []
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text.strip())
    # Also read tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            text.append(" | ".join(cells))
    return "\n".join(text)

base = r"C:\Odoo Study\My learnings\Rohan_Documentation"

print("=" * 80)
print("DOC 1: Multiple_Branches.docx")
print("=" * 80)
content1 = read_docx(os.path.join(base, "Multiple_Branches.docx"))
print(content1[:5000])
print("\n... (truncated)" if len(content1) > 5000 else "")

print("\n\n" + "=" * 80)
print("DOC 2: Multi_Company_Branches_Warehouse_Learnings.docx")
print("=" * 80)
content2 = read_docx(os.path.join(base, "Multi_Company_Branches_Warehouse_Learnings.docx"))
print(content2[:5000])
print("\n... (truncated)" if len(content2) > 5000 else "")

# ── Check client-cient database ──
print("\n\n" + "=" * 80)
print("CLIENT-CIENT DATABASE SETUP")
print("=" * 80)

URL = "https://client-cient.odoo.com"
DB = "client-cient"
USERNAME = "rohan.raj@infintor.com"
PASSWORD = "Rohanraj@1"

common = xmlrpc.client.ServerProxy(f"{URL}/xmlrpc/2/common")
uid = common.authenticate(DB, USERNAME, PASSWORD, {})
models = xmlrpc.client.ServerProxy(f"{URL}/xmlrpc/2/object")

def search_read(model, domain, fields, limit=50):
    return models.execute_kw(DB, uid, PASSWORD, model, 'search_read',
                             [domain], {'fields': fields, 'limit': limit})

# Companies
print("\n--- COMPANIES ---")
companies = search_read('res.company', [], ['name', 'parent_id', 'street', 'city', 'state_id', 'country_id'])
for c in companies:
    parent = c['parent_id'][1] if c['parent_id'] else 'None'
    state = c['state_id'][1] if c['state_id'] else ''
    print(f"  {c['name']} | Parent: {parent} | City: {c['city'] or ''} | State: {state}")

# Warehouses
print("\n--- WAREHOUSES ---")
warehouses = search_read('stock.warehouse', [], ['name', 'code', 'company_id', 'partner_id'])
for w in warehouses:
    company = w['company_id'][1] if w['company_id'] else ''
    print(f"  {w['name']} ({w['code']}) | Company: {company}")

# Locations (internal only, key ones)
print("\n--- KEY STOCK LOCATIONS ---")
locations = search_read('stock.location', [['usage', '=', 'internal']], ['complete_name', 'company_id', 'warehouse_id'])
for loc in locations:
    company = loc['company_id'][1] if loc['company_id'] else ''
    wh = loc['warehouse_id'][1] if loc['warehouse_id'] else ''
    print(f"  {loc['complete_name']} | WH: {wh} | Company: {company}")

# Inter-company/inter-warehouse transfers (resupply routes)
print("\n--- RESUPPLY ROUTES ---")
try:
    routes = search_read('stock.route', [], ['name', 'company_id', 'supplied_wh_id', 'supplier_wh_id'])
    for r in routes:
        company = r['company_id'][1] if r['company_id'] else 'All'
        supplied = r.get('supplied_wh_id', [False, ''])[1] if r.get('supplied_wh_id') else ''
        supplier = r.get('supplier_wh_id', [False, ''])[1] if r.get('supplier_wh_id') else ''
        print(f"  {r['name']} | Company: {company} | From: {supplier} → To: {supplied}")
except:
    print("  Could not read routes")

# Products count
print("\n--- PRODUCTS ---")
products = search_read('product.template', [['type', '!=', 'service']], ['name', 'type', 'tracking', 'company_id'], limit=30)
print(f"  Total storable/consumable products (first 30):")
for p in products:
    company = p['company_id'][1] if p['company_id'] else 'All companies'
    print(f"    {p['name']} | Type: {p['type']} | Tracking: {p['tracking']} | Company: {company}")

# Check inter-company rules
print("\n--- INTER-COMPANY RULES ---")
try:
    rules = search_read('res.inter.company.rule', [], ['name', 'company_id', 'rule_type'])
    if rules:
        for r in rules:
            print(f"  {r.get('name', 'N/A')} | Company: {r['company_id'][1]} | Type: {r.get('rule_type', 'N/A')}")
    else:
        print("  No inter-company rules configured")
except:
    print("  Model not available or no rules")

# Check purchase/sale orders
print("\n--- RECENT PURCHASE ORDERS ---")
pos = search_read('purchase.order', [], ['name', 'partner_id', 'company_id', 'state', 'date_order'], limit=10)
for po in pos:
    print(f"  {po['name']} | Vendor: {po['partner_id'][1]} | Company: {po['company_id'][1]} | State: {po['state']}")

print("\n--- RECENT SALE ORDERS ---")
sos = search_read('sale.order', [], ['name', 'partner_id', 'company_id', 'state', 'warehouse_id'], limit=10)
for so in sos:
    wh = so['warehouse_id'][1] if so['warehouse_id'] else ''
    print(f"  {so['name']} | Customer: {so['partner_id'][1]} | Company: {so['company_id'][1]} | WH: {wh} | State: {so['state']}")

# Internal transfers
print("\n--- INTERNAL TRANSFERS ---")
transfers = search_read('stock.picking', [['picking_type_code', '=', 'internal']], 
    ['name', 'origin', 'company_id', 'location_id', 'location_dest_id', 'state'], limit=15)
for t in transfers:
    src = t['location_id'][1] if t['location_id'] else ''
    dest = t['location_dest_id'][1] if t['location_dest_id'] else ''
    print(f"  {t['name']} | {src} → {dest} | Company: {t['company_id'][1]} | State: {t['state']}")

# Users
print("\n--- USERS WITH COMPANY ACCESS ---")
users = search_read('res.users', [['share', '=', False]], ['name', 'company_id', 'company_ids'], limit=10)
for u in users:
    main = u['company_id'][1] if u['company_id'] else ''
    allowed = u.get('company_ids', [])
    print(f"  {u['name']} | Main: {main} | Allowed companies: {allowed}")

print("\n=== DONE ===")
