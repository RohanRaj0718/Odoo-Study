"""Convert CLIENT_OPERATIONS_GUIDE.md to a clean, professional Word document
following the style of Profit and Loss Reporting.docx and Account Grouping Doc.docx"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Paragraph spacing
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Title
title = doc.add_heading('How to Manage Sales, Purchases, Internal Transfers and Reporting Across Your Warehouses and Companies in Odoo 19', level=1)
title.runs[0].font.size = Pt(16)
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

# Author & Date
meta = doc.add_paragraph('Rohan Raj | Feb 23, 2026')
meta.runs[0].font.size = Pt(10)
meta.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
meta.runs[0].font.italic = True

# Intro
doc.add_paragraph(
    'This document explains how to use the existing setup in Odoo 19 to create sales orders, purchase orders, '
    'transfer stock between warehouses automatically, and view financial reports for each company and branch. '
    'The setup currently includes two main companies and two branches, with four warehouses under the parent company. '
    'All steps described below are based on the live configuration at psi-122test.odoo.com.'
)

# --- SECTION: Your Current Setup ---
doc.add_heading('Your Current Setup at a Glance', level=2)

doc.add_paragraph(
    'The database contains two independent companies and two branches within the parent company.'
)

p = doc.add_paragraph()
r = p.add_run('Parappattu Group')
r.bold = True
p.add_run(
    ' is the parent company (GSTIN: 32AAOFP3684P1ZH). It has four warehouses:'
)

items = [
    ('Parappattu Group (WH)', 'Main headquarters'),
    ('Near Home GF (NH GF)', 'Ground floor branch store'),
    ('Near Home FF (NH FF)', 'First floor branch store'),
    ('Factory Building (FB)', 'Factory and storage location'),
]
for name, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(name)
    r.bold = True
    p.add_run(f' — {desc}')

doc.add_paragraph(
    'Parappattu Group also has two child companies (branches): Georgeon Furniture and PSQUARE INTERIOR. '
    'These branches share the same GSTIN as the parent. Since all four warehouses belong to Parappattu Group, '
    'all sales, purchases and transfers happen under the Parappattu Group books.'
)

p = doc.add_paragraph()
r = p.add_run('PSQUARE INTERIOR FURNISHING')
r.bold = True
p.add_run(
    ' is a separate independent company with its own GSTIN (32AAKCP9897R1Z4) and its own chart of accounts. '
    'It currently has no warehouse and needs one to be created before it can handle inventory.'
)

doc.add_paragraph(
    'The single user Georgey has access to all four companies from the company selector in the top right corner of the screen.'
)

# --- PART 1: Sales Orders ---
doc.add_heading('Part 1 — Creating a Sales Order from a Specific Warehouse', level=2)

doc.add_paragraph(
    'When a customer visits one of the branch stores or places an order, a sales order is created and '
    'the delivery warehouse is selected to ensure goods are shipped from the correct location.'
)

doc.add_paragraph(
    'To create a sales order from a specific warehouse, open the Sales application and click New to create '
    'a new quotation. Select the customer, then add the products and quantities required.'
)

doc.add_paragraph(
    'Before confirming the order, look for the Warehouse field on the sales order form. By default this is '
    'set to Parappattu Group (WH). To ship from a different location, click the Warehouse field and select '
    'the appropriate warehouse:'
)

wh_items = [
    ('Near Home GF (NH GF)', 'for sales from the ground floor branch'),
    ('Near Home FF (NH FF)', 'for sales from the first floor branch'),
    ('Factory Building (FB)', 'for goods shipped directly from the factory'),
]
for name, desc in wh_items:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(name)
    r.bold = True
    p.add_run(f' — {desc}')

# Navigation path
nav = doc.add_paragraph()
nav.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = nav.add_run('Sales App → New → Select Customer → Add Products → Change Warehouse → Confirm')
r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.size = Pt(10)

doc.add_paragraph(
    'Once the order is confirmed, Odoo automatically creates a delivery order from the selected warehouse. '
    'The invoice is generated under Parappattu Group since all warehouses belong to this company.'
)

doc.add_paragraph(
    'To track which branch made the sale, use the Sales Team field or a tag on the sales order. '
    'This is useful when reviewing branch-level performance in reports later.'
)

# --- PART 2: Purchase Orders ---
doc.add_heading('Part 2 — Creating a Purchase Order and Receiving Goods into a Specific Warehouse', level=2)

doc.add_paragraph(
    'When goods are purchased from a vendor, they can be received directly into any of the four warehouses. '
    'This avoids the need to first receive everything at headquarters and then transfer to branches.'
)

doc.add_paragraph(
    'To create a purchase order, open the Purchase application and click New. Select the vendor, '
    'then add the products and quantities.'
)

doc.add_paragraph(
    'On the purchase order form, locate the Deliver To field. This field determines which warehouse will '
    'receive the incoming goods. Click on this field and select the destination warehouse. For example, '
    'to receive goods directly at the factory, select Factory Building (FB). To receive at a branch store, '
    'select Near Home GF (NH GF) or Near Home FF (NH FF).'
)

nav = doc.add_paragraph()
nav.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = nav.add_run('Purchase App → New → Select Vendor → Add Products → Change Deliver To → Confirm Order')
r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.size = Pt(10)

doc.add_paragraph(
    'After the order is confirmed and the vendor ships the goods, a receipt is created at the selected warehouse. '
    'When the goods arrive, open the receipt and click Validate to confirm that the products have been received. '
    'All purchase bills are recorded under Parappattu Group\'s accounting.'
)

# --- PART 3: Internal Transfers ---
doc.add_heading('Part 3 — Internal Transfers Between Warehouses', level=2)

doc.add_paragraph(
    'When stock needs to move from one warehouse to another within Parappattu Group, an internal transfer is created. '
    'Since all four warehouses belong to the same company, this is a simple stock movement with no tax implications '
    'and no invoicing required.'
)

doc.add_heading('Method 1 — Manual Internal Transfer', level=3)

doc.add_paragraph(
    'To create a manual transfer, navigate to the Inventory application. From the top menu, go to Operations '
    'and select Transfers. Click New to create a new transfer.'
)

nav = doc.add_paragraph()
nav.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = nav.add_run('Inventory App → Operations → Transfers → New')
r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.size = Pt(10)

doc.add_paragraph(
    'In the transfer form, set the Operation Type to Internal Transfer. In the Source Location field, select '
    'the warehouse location from which the goods are being sent (for example WH/Stock for the main warehouse). '
    'In the Destination Location field, select the warehouse location that will receive the goods '
    '(for example NH GF/Stock for the Near Home GF branch).'
)

doc.add_paragraph(
    'Add the products and quantities to be transferred. Click Mark as To Do, then when the goods are physically '
    'moved, click Validate to complete the transfer. The stock levels at both locations are updated immediately.'
)

doc.add_heading('Method 2 — Automatic Resupply Between Warehouses', level=3)

doc.add_paragraph(
    'Instead of creating manual transfers each time, Odoo can be configured to automatically move stock from a '
    'supply warehouse to a branch warehouse whenever stock runs low. This is done using resupply routes and reordering rules.'
)

# Step 1
p = doc.add_paragraph()
r = p.add_run('Step 1 — Enable Multi-Step Routes')
r.bold = True

doc.add_paragraph(
    'Before setting up resupply, Multi-Step Routes must be enabled in the inventory settings.'
)

nav = doc.add_paragraph()
nav.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = nav.add_run('Inventory App → Configuration → Settings → Warehouse section → Enable Multi-Step Routes → Save')
r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.size = Pt(10)

# Step 2
p = doc.add_paragraph()
r = p.add_run('Step 2 — Configure Resupply on the Branch Warehouse')
r.bold = True

doc.add_paragraph(
    'Open the warehouse that needs automatic resupply. For example, to have Near Home GF automatically pull '
    'stock from the main warehouse:'
)

nav = doc.add_paragraph()
nav.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = nav.add_run('Inventory App → Configuration → Warehouses → Select Near Home GF (NH GF)')
r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.size = Pt(10)

doc.add_paragraph(
    'In the warehouse settings, find the Resupply From field. Check the box next to Parappattu Group (WH). '
    'Save the changes. This creates a new route called "Near Home GF: Supply Product from Parappattu Group" '
    'which handles the automatic transfer.'
)

doc.add_paragraph(
    'Repeat this for any other warehouse that should be automatically resupplied. For example, configure '
    'Near Home FF to resupply from Parappattu Group, or from Factory Building.'
)

# Step 3
p = doc.add_paragraph()
r = p.add_run('Step 3 — Set Up Reordering Rules')
r.bold = True

doc.add_paragraph(
    'After enabling resupply, reordering rules tell Odoo when to trigger the automatic transfer. For each product '
    'at each branch warehouse, a reordering rule defines the minimum and maximum stock levels.'
)

nav = doc.add_paragraph()
nav.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = nav.add_run('Inventory App → Operations → Replenishment → New')
r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.size = Pt(10)

doc.add_paragraph(
    'Select the product, choose the branch warehouse location (for example NH GF/Stock), set the minimum quantity '
    '(the level below which stock should be replenished) and the maximum quantity (the target level after replenishment). '
    'Set the preferred route to the resupply route created in Step 2.'
)

doc.add_paragraph(
    'When the stock at the branch falls below the minimum, Odoo automatically creates an internal transfer from '
    'the supply warehouse, moving enough stock to reach the maximum quantity. The scheduler runs daily or can be '
    'triggered manually from Inventory → Operations → Run Scheduler.'
)

# --- PART 4: PSQUARE INTERIOR FURNISHING ---
doc.add_heading('Part 4 — Operations for PSQUARE INTERIOR FURNISHING (Separate Company)', level=2)

doc.add_paragraph(
    'PSQUARE INTERIOR FURNISHING is a separate legal entity with a different GSTIN. It cannot use Parappattu Group\'s '
    'warehouses. Any exchange of goods between these two companies requires a proper sale and purchase transaction with GST.'
)

p = doc.add_paragraph()
r = p.add_run('Step 1 — Create a Warehouse for PSQUARE INTERIOR FURNISHING')
r.bold = True

doc.add_paragraph(
    'Switch to the PSQUARE INTERIOR FURNISHING company using the company selector in the top right corner. '
    'Then navigate to warehouse configuration.'
)

nav = doc.add_paragraph()
nav.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = nav.add_run('Company Selector → PSQUARE INTERIOR FURNISHING → Inventory App → Configuration → Warehouses → New')
r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.size = Pt(10)

doc.add_paragraph(
    'Enter the warehouse name (for example "PSQUARE Warehouse"), a short code (for example "PSQ"), and ensure '
    'the company is set to PSQUARE INTERIOR FURNISHING. Save the warehouse.'
)

p = doc.add_paragraph()
r = p.add_run('Step 2 — How Stock Moves Between the Two Companies')
r.bold = True

doc.add_paragraph(
    'When PSQUARE INTERIOR FURNISHING needs products from Parappattu Group (or vice versa), the process works '
    'like a normal business transaction between two separate entities:'
)

steps_inter = [
    'Under Parappattu Group, create a sales order with PSQUARE INTERIOR FURNISHING as the customer. Confirm the order, deliver the goods, and create the invoice with applicable GST.',
    'Under PSQUARE INTERIOR FURNISHING, create a purchase order with Parappattu Group as the vendor. Confirm the order, receive the goods, and record the vendor bill.',
]
for i, step in enumerate(steps_inter, 1):
    p = doc.add_paragraph(style='List Number')
    p.text = step

doc.add_paragraph(
    'This ensures proper GST compliance since both companies have different GSTINs and the transaction is legally '
    'a sale and purchase between two entities.'
)

# --- PART 5: Financial Reporting ---
doc.add_heading('Part 5 — Viewing Financial Reports by Company', level=2)

doc.add_paragraph(
    'Odoo provides separate financial reports for each company. By switching companies using the company selector, '
    'the reports automatically show data for the selected company only.'
)

doc.add_heading('Viewing Profit and Loss for Parappattu Group', level=3)

doc.add_paragraph(
    'To check whether Parappattu Group is profitable for a specific period, switch to Parappattu Group using the '
    'company selector, then open the reporting section.'
)

nav = doc.add_paragraph()
nav.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = nav.add_run('Company Selector → Parappattu Group → Accounting App → Reporting → Profit and Loss')
r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.size = Pt(10)

doc.add_paragraph(
    'The Profit and Loss report displays revenue earned during the period, followed by the cost of revenue to '
    'calculate gross profit. Operating expenses are subtracted to show operating income or loss. Other income and '
    'expenses are then applied to show the net profit. To change the period, click the date button at the top of '
    'the report and select the required month, quarter, year, or custom date range.'
)

doc.add_heading('Viewing Profit and Loss for PSQUARE INTERIOR FURNISHING', level=3)

doc.add_paragraph(
    'Switch to PSQUARE INTERIOR FURNISHING using the company selector, then open the same report.'
)

nav = doc.add_paragraph()
nav.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = nav.add_run('Company Selector → PSQUARE INTERIOR FURNISHING → Accounting App → Reporting → Profit and Loss')
r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.size = Pt(10)

doc.add_paragraph(
    'The report now shows only the financial data for PSQUARE INTERIOR FURNISHING, as it has its own separate chart '
    'of accounts, journals and transactions.'
)

doc.add_heading('Viewing Reports for Branches (Georgeon Furniture and PSQUARE INTERIOR)', level=3)

doc.add_paragraph(
    'Georgeon Furniture and PSQUARE INTERIOR are child companies (branches) of Parappattu Group. They share the same '
    'GSTIN as the parent. Currently, these branches do not have their own chart of accounts or warehouses. '
    'All operations run through Parappattu Group.'
)

doc.add_paragraph(
    'To track branch-level performance within Parappattu Group, there are two recommended approaches.'
)

# Approach 1
p = doc.add_paragraph()
r = p.add_run('Approach 1 — Using Analytic Accounts')
r.bold = True

doc.add_paragraph(
    'Analytic accounting allows tagging each transaction (sale, purchase, expense) with a branch identifier '
    'without creating separate company books. First, enable analytic accounting in the settings.'
)

nav = doc.add_paragraph()
nav.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = nav.add_run('Accounting App → Configuration → Settings → Analytics section → Enable Analytic Accounting → Save')
r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.size = Pt(10)

doc.add_paragraph(
    'Then create an analytic plan called "Branches" and add analytic accounts for each branch '
    '(Georgeon Furniture, PSQUARE INTERIOR, Near Home GF, Near Home FF, Factory Building).'
)

nav = doc.add_paragraph()
nav.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = nav.add_run('Accounting App → Configuration → Analytic Plans → New → Name: Branches')
r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.size = Pt(10)

doc.add_paragraph(
    'When creating sales orders, purchase orders, or journal entries, select the appropriate analytic account '
    'to tag the transaction to the correct branch. To view branch performance, open the Profit and Loss report '
    'and filter by analytic account.'
)

# Approach 2
p = doc.add_paragraph()
r = p.add_run('Approach 2 — Using Sales Teams')
r.bold = True

doc.add_paragraph(
    'For tracking sales performance by branch, create a separate sales team for each branch.'
)

nav = doc.add_paragraph()
nav.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = nav.add_run('Sales App → Configuration → Sales Teams → New')
r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.size = Pt(10)

doc.add_paragraph(
    'Create teams such as "Georgeon Sales", "PSQUARE Sales", "Near Home GF Sales" and so on. '
    'When creating sales orders, assign the appropriate sales team. To view performance by branch, '
    'open Sales App → Reporting → Sales → Group By → Sales Team.'
)

# --- PART 6: Inventory Reports ---
doc.add_heading('Part 6 — Viewing Inventory Reports by Warehouse', level=2)

doc.add_heading('Stock by Warehouse', level=3)

doc.add_paragraph(
    'To see what stock is available at each warehouse location:'
)

nav = doc.add_paragraph()
nav.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = nav.add_run('Inventory App → Reporting → Inventory Report → Group By → Warehouse')
r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.size = Pt(10)

doc.add_paragraph(
    'This shows the on-hand quantity, forecasted quantity and available quantity at each warehouse separately. '
    'To check a specific product\'s availability across all warehouses, open the product form and click the '
    'On Hand smart button.'
)

doc.add_heading('Stock Moves History', level=3)

nav = doc.add_paragraph()
nav.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = nav.add_run('Inventory App → Reporting → Stock Moves')
r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.size = Pt(10)

doc.add_paragraph(
    'Filter by source location or destination location to see transfers involving a specific warehouse. '
    'For example, filter the source location to WH/Stock to see all items sent out from the main warehouse.'
)

# --- PART 7: Quick Reference ---
doc.add_heading('Part 7 — Quick Reference Summary', level=2)

# Create table
table = doc.add_table(rows=9, cols=2)
table.style = 'Light Shading Accent 1'

headers = ['Operation', 'How To']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

ref_data = [
    ('Sale from a branch', 'Sales App → New → Set Warehouse to the branch → Confirm'),
    ('Purchase for a branch', 'Purchase App → New → Set Deliver To to the branch warehouse → Confirm'),
    ('Internal transfer', 'Inventory App → Operations → Transfers → New → Set source and destination'),
    ('Automatic resupply', 'Inventory → Configuration → Warehouses → Enable Resupply From → Create reordering rules'),
    ('Company Profit and Loss', 'Switch company → Accounting → Reporting → Profit and Loss → Select period'),
    ('Branch-level profitability', 'Enable Analytic Accounting → Tag transactions → Filter P&L by analytic'),
    ('Stock at a warehouse', 'Inventory App → Reporting → Inventory Report → Group By Warehouse'),
    ('Inter-company transfer', 'Create SO in sending company → Create PO in receiving company → Deliver and invoice'),
]

for i, (op, how) in enumerate(ref_data, 1):
    table.rows[i].cells[0].text = op
    table.rows[i].cells[1].text = how

# Closing note
doc.add_paragraph('')
doc.add_paragraph(
    'This document covers the main daily operations for the current setup. For any additional configuration '
    'such as adding new warehouses, setting up user access restrictions per branch, or enabling E-Waybill for '
    'inter-state transfers, refer to the detailed technical setup guide or contact the implementation team.'
)

# Save
doc.save('CLIENT_OPERATIONS_GUIDE.docx')
print("Document saved: CLIENT_OPERATIONS_GUIDE.docx")
